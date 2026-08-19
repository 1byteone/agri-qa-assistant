"""Safe document parsing and agriculture-domain preflight for RAG ingestion."""
from __future__ import annotations

import hashlib
import html
import io
import json
import re
from pathlib import Path
from typing import Any, Dict

from domain_guard import AGRICULTURE_TERMS, NON_AGRICULTURE_TERMS

MAX_UPLOAD_BYTES = 15 * 1024 * 1024
MAX_TEXT_CHARS = 500_000
MIN_TEXT_CHARS = 40

SUPPORTED_EXTENSIONS = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".csv": "text/csv",
    ".html": "text/html",
    ".htm": "text/html",
    ".json": "application/json",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
}


class DocumentIngestionError(ValueError):
    """A user-correctable document ingestion failure."""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        document = Document(io.BytesIO(data))
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)
    except ImportError as exc:  # pragma: no cover
        raise DocumentIngestionError("服务器未安装 DOCX 解析依赖") from exc
    except Exception as exc:
        raise DocumentIngestionError("DOCX 文件解析失败，请确认文件未损坏") from exc


def _extract_pdf(data: bytes) -> str:
    try:
        import fitz
        with fitz.open(stream=data, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)
    except ImportError as exc:  # pragma: no cover
        raise DocumentIngestionError("服务器未安装 PDF 解析依赖") from exc
    except Exception as exc:
        raise DocumentIngestionError("PDF 文件解析失败，请确认文件未损坏或包含可提取文本") from exc


def _extract_html(data: bytes) -> str:
    raw = _decode_text(data)
    try:
        from bs4 import BeautifulSoup
        return BeautifulSoup(raw, "html.parser").get_text("\n")
    except ImportError:  # pragma: no cover
        return re.sub(r"<[^>]+>", " ", raw)


def _extract_content(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".docx":
        return _extract_docx(data)
    if extension == ".pdf":
        return _extract_pdf(data)
    if extension in {".html", ".htm"}:
        return _extract_html(data)
    if extension == ".json":
        try:
            return json.dumps(json.loads(_decode_text(data)), ensure_ascii=False, indent=2)
        except json.JSONDecodeError as exc:
            raise DocumentIngestionError("JSON 文件格式无效") from exc
    return _decode_text(data)


def _normalise_text(text: str) -> str:
    text = html.unescape(text or "").replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def analyse_agriculture_document(text: str) -> Dict[str, Any]:
    """Return a conservative and explainable domain decision."""
    compact = re.sub(r"\s+", " ", text.lower())
    matched = sorted({term for term in AGRICULTURE_TERMS if term in compact}, key=len, reverse=True)
    non_agriculture = sorted({term for term in NON_AGRICULTURE_TERMS if term in compact}, key=len, reverse=True)
    strong_matches = [term for term in matched if term not in {"农业", "农产品"}]
    negative_scope = bool(re.search(r"(?:没有|不含|不包含|不涉及|非|不属于).{0,8}(?:农业|农技|作物)", compact))
    code_terms = {"java", "python", "javascript", "typescript", "代码", "编程", "程序", "开发"}
    mixed_code = bool(code_terms.intersection(non_agriculture)) and len(strong_matches) < 2
    eligible = bool(matched) and not (non_agriculture and not matched) and not negative_scope and not mixed_code
    confidence = min(1.0, 0.35 + 0.1 * min(len(matched), 6)) if eligible else 0.0
    if not eligible:
        reason = "未识别到足够的农业主题词，文件不会进入农业知识库"
    elif len(matched) == 1:
        reason = "识别到农业主题，但建议人工核对文档内容后再入库"
    else:
        reason = "已识别到农业主题，可在确认后进入农业知识库"
    return {
        "eligible": eligible,
        "confidence": round(confidence, 2),
        "matched_terms": matched[:12],
        "non_agriculture_terms": non_agriculture[:8],
        "reason": reason,
    }


def parse_document(filename: str | None, content_type: str | None, data: bytes) -> Dict[str, Any]:
    """Parse a bounded upload and return analysis without mutating the vector DB."""
    safe_name = Path(filename or "").name
    extension = Path(safe_name).suffix.lower()
    if not safe_name or extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentIngestionError(f"暂不支持该文件格式，仅支持：{supported}")
    if not data:
        raise DocumentIngestionError("文件为空，无法解析")
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentIngestionError(f"文件超过 {MAX_UPLOAD_BYTES // (1024 * 1024)}MB 大小限制")
    text = _normalise_text(_extract_content(safe_name, data))
    if len(text) < MIN_TEXT_CHARS:
        raise DocumentIngestionError(f"可解析文本过少（至少需要 {MIN_TEXT_CHARS} 个字符）")
    if len(text) > MAX_TEXT_CHARS:
        raise DocumentIngestionError(f"可解析文本超过 {MAX_TEXT_CHARS:,} 个字符限制")
    analysis = analyse_agriculture_document(text)
    return {
        "filename": safe_name,
        "extension": extension,
        "content_type": content_type or SUPPORTED_EXTENSIONS[extension],
        "bytes": len(data),
        "characters": len(text),
        "estimated_chunks": max(1, (len(text) + 799) // 800),
        "content_hash": hashlib.sha256(data).hexdigest(),
        "preview": text[:600],
        "text": text,
        **analysis,
    }


def public_analysis(parsed: Dict[str, Any]) -> Dict[str, Any]:
    """Remove full extracted text before returning an analysis response."""
    return {key: value for key, value in parsed.items() if key != "text"}
